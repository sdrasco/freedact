#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
freedact.py — Self‑contained, offline PII redactor for PDF/DOCX/DOC (Python 3.9+)

Summary
-------
Redacts personally identifiable information from documents **entirely offline** and
writes:
  1) <input>_redacted.docx (always attempted)
  2) <input>_redacted.pdf (when --pdf is passed)
  3) <input>_redaction_key.json (mapping placeholders → original names/aliases, unless --no-keep-key)
  4) A concise console summary of redaction counts by category.

Deterministic placeholder assignment (John Doe 1, John Doe 2, …) follows order of first
appearance. Heuristics are simple by design, prioritizing safety and portability over
layout fidelity. Works on macOS and is cross‑platform where dependencies are available.

Core Features
-------------
• Persons:
  - Detect sequences of ≥2 capitalized tokens with optional titles (Dr./Mr./Mrs./Ms./Prof.)
    and optional middle initials (e.g., "J. Smith", "Jane A. Smith").
  - Skip obvious organizations (university, college, credit union, bank, trust, llc, inc,
    corp, company, foundation, department, committee, society).
  - Link “Hereinafter ‘Alias’” style nicknames to the corresponding full name.
  - Preserve possessives (e.g., “Smith’s car” → “John Doe 1’s car”).
  - Optional --include-allcaps to consider ALL-CAPS candidate names (default off).
  - Optional --mask-mode to use [REDACTED] instead of John Doe N in body text.

• Addresses → [REDACTED ADDRESS] (whole-line redaction):
  - Lines like “<number> <street> St/Rd/Ave/…”, “City, ST 12345” (ZIP+4 supported),
    common UK postcodes (e.g., “KY16 8QP”), P.O. Box lines.
  - If a street line is redacted, redact immediately following Apt/Apartment/Unit line.

• Account Names → [REDACTED ACCOUNT NAME]
  - Built-in extensible list (e.g., “retirement account”, “brokerage account”, “checking
    account”, “money market”, “credit card”, “current account”, “Citibank”, “Fidelity”,
    “TIAA‑CREF”, “CalPERS”, “Santander”, common “credit union”).
  - Add more at runtime with --account-term "Custom Term" (repeatable).

• Account Numbers/IDs → [REDACTED ACCOUNT NUMBER]
  - Patterns include: “account #…”, IBAN-like, 7+ contiguous digits, card-like
    “#### #### #### ####”.
  - Optional --strict-ids to also redact VIN-like IDs (A‑Z0‑9, 11–17 chars, excluding I/O/Q).

Inputs
------
• .pdf, .docx, .doc
• PDF (text-based): try `pdfplumber` first, else `pypdf` (PyPDF2/pypdf).
• PDF (scanned): optional OCR via `pytesseract` + `pdf2image` when --ocr is passed.
  (Requires system installs of Tesseract and Poppler.)
• DOCX: `python-docx`.
• DOC (legacy): try `textract`, else `antiword`. If both unavailable, instruct to convert to .docx.

Writers
-------
• DOCX: `python-docx`. Writes redacted paragraphs; appends “Anonymization Key (Persons)”
  page (unless --no-keep-key).
• PDF: `reportlab`, A4 portrait, Courier, simple word-wrap; final page prints the same key
  (unless --no-keep-key).

Privacy
-------
• No network calls, telemetry, or external services. Works entirely offline on local files.

Installation (offline-friendly)
-------------------------------
Python packages (create/activate a venv if desired):
    pip install --upgrade pip
    pip install python-docx reportlab unidecode
    pip install pdfplumber  # or: pip install pypdf
    # Optional for .doc:
    pip install textract
    # Optional OCR for scanned PDFs (used only with --ocr):
    pip install pytesseract pdf2image

macOS (Homebrew) system tools (only if you need them):
    brew install tesseract poppler   # OCR support for --ocr (tesseract & pdftoppm)
    brew install antiword            # Optional .doc fallback

Windows (system tools if you need OCR):
  - Install Tesseract: https://tesseract-ocr.github.io/tessdoc/Installation.html
  - Install Poppler:   https://blog.alivate.com.au/poppler-windows/

Usage
-----
    python freedact.py input.pdf --pdf --ocr
    python freedact.py input.docx --strict-ids --account-term "Acme Bank"
    python freedact.py input.doc --pdf
    python freedact.py input.pdf --dry-run
    python freedact.py --self-test

Run `python freedact.py --help` for full CLI details.

License
-------
This script is provided "as is" without warranty of any kind. Review outputs before use.

"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import sys
import textwrap
from collections import OrderedDict, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Iterable

# --- Required runtime check for 'unidecode' (hard requirement for normalization) ---
try:
    from unidecode import unidecode
except Exception:
    sys.stderr.write(
        "[ERROR] Missing required dependency 'unidecode'.\n"
        "Install it with:\n    pip install unidecode\n"
    )
    sys.exit(1)


# ----------------------------
# Utility dataclasses & types
# ----------------------------

@dataclass
class RedactionResult:
    text: str
    counts: Dict[str, int]
    placeholder_map: "OrderedDict[str, Dict[str, List[str]]]"  # {placeholder: {'canonical': str, 'aliases': [..]}}


# ----------------------------
# CLI / Help Text
# ----------------------------

HELP_EPILOG = """
Examples:
  python freedact.py input.pdf --pdf --ocr
  python freedact.py input.docx --strict-ids --account-term "Acme Bank"
  python freedact.py input.doc --pdf
  python freedact.py input.pdf --dry-run

Dependencies:
  Required (core): python-docx, reportlab, unidecode, and either pdfplumber or pypdf (for PDF text).
  Optional (.doc): textract and/or antiword (Homebrew: `brew install antiword`).
  Optional OCR: pytesseract, pdf2image (Homebrew: `brew install tesseract poppler`).

Notes:
  • The DOCX output is always attempted. If 'python-docx' is missing, you'll get actionable install guidance.
  • The PDF output is written only when --pdf is passed and 'reportlab' is installed.
  • For PDFs: if text extraction yields nothing and --ocr is provided, OCR is used when pytesseract/pdf2image are installed.
  • All processing is deterministic for a given input and flag set.
  • This tool prioritizes privacy: no network calls, telemetry, or external services.
""".strip()


def build_arg_parser() -> argparse.ArgumentParser:
    desc = (
        "Offline PII redactor for PDF/DOCX/DOC.\n\n"
        "Installs (quick start):\n"
        "  pip install python-docx reportlab unidecode pdfplumber\n"
        "  # or: pip install pypdf\n"
        "Optional:\n"
        "  pip install textract              # .doc support\n"
        "  pip install pytesseract pdf2image # OCR for scanned PDFs (use with --ocr)\n"
        "Homebrew (macOS, optional):\n"
        "  brew install tesseract poppler    # OCR prerequisites for --ocr\n"
        "  brew install antiword             # .doc fallback\n"
    )
    parser = argparse.ArgumentParser(
        prog="freedact.py",
        description=desc,
        epilog=HELP_EPILOG,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", nargs="?", help="Path to input file (.pdf, .docx, .doc)")
    parser.add_argument("--pdf", action="store_true", help="Also write <input>_redacted.pdf (uses 'reportlab')")
    parser.add_argument("--ocr", action="store_true", help="Enable OCR for PDFs when text is not extractable")
    parser.add_argument("--strict-ids", action="store_true", help="Also redact VIN-like IDs (A-Z0-9, 11–17 chars, no I/O/Q)")
    parser.add_argument("--include-allcaps", action="store_true", help="Treat ALL-CAPS as candidate person names")
    parser.add_argument("--account-term", action="append", default=[], help="Add custom account label/brand to redact (repeatable)")
    parser.add_argument("--mask-mode", action="true_false", nargs="?", const=True, default=False,
                        help="Use [REDACTED] for names instead of placeholders (default: off)")
    parser.add_argument("--keep-key", dest="keep_key", action="true_false", nargs="?", const=True, default=True,
                        help="Write JSON mapping and append key page (default: on). Disable with --no-keep-key")
    parser.add_argument("--no-keep-key", dest="keep_key", action="store_false", help=argparse.SUPPRESS)
    parser.add_argument("--dry-run", action="store_true", help="Print what would be redacted (counts); do not write files")
    parser.add_argument("--self-test", action="store_true", help="Run built-in acceptance tests and exit")
    return parser


# argparse doesn't have a built-in 'true_false' action; implement a tiny shim:
class _TrueFalseAction(argparse.Action):
    def __call__(self, parser, ns, values, option_string=None):
        setattr(ns, self.dest, values if isinstance(values, bool) else True)


# Register shim
argparse._ActionsContainer.register("true_false", _TrueFalseAction)


# ----------------------------
# Reading inputs (offline)
# ----------------------------

def read_input_text(path: Path, use_ocr: bool) -> str:
    """
    Read the input file and return a single text string.
    Uses only offline libraries. Gracefully degrades with guidance if optional deps are missing.

    PDF: pdfplumber preferred, else pypdf. If both unavailable and --ocr is set, tries OCR via
         pdf2image + pytesseract (requires system installs).
    DOCX: python-docx
    DOC:  textract, else antiword (system), else instruct conversion.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf_text(path)
        if text and text.strip():
            return text
        if use_ocr:
            return _read_pdf_via_ocr(path)
        msg = (
            "[ERROR] No text extracted from PDF. Try installing 'pdfplumber' or 'pypdf'.\n"
            "       For scanned PDFs, re-run with --ocr and install:\n"
            "         pip install pytesseract pdf2image\n"
            "         brew install tesseract poppler\n"
        )
        sys.stderr.write(msg)
        sys.exit(2)

    elif suffix == ".docx":
        try:
            import docx  # type: ignore
        except Exception:
            sys.stderr.write(
                "[ERROR] Missing 'python-docx' for reading .docx.\n"
                "Install it with:\n    pip install python-docx\n"
            )
            sys.exit(2)
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    elif suffix == ".doc":
        # Try textract first
        try:
            import textract  # type: ignore
            data = textract.process(str(path))
            return data.decode("utf-8", errors="ignore")
        except Exception:
            # Try antiword system tool
            try:
                import subprocess
                proc = subprocess.run(
                    ["antiword", str(path)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=False,
                )
                if proc.returncode == 0:
                    return proc.stdout.decode("utf-8", errors="ignore")
                else:
                    raise RuntimeError(proc.stderr.decode("utf-8", errors="ignore"))
            except Exception:
                sys.stderr.write(
                    "[ERROR] Cannot read .doc. Options:\n"
                    "  • Install 'textract': pip install textract\n"
                    "  • Or install antiword (macOS): brew install antiword\n"
                    "  • Or convert the .doc to .docx and retry.\n"
                )
                sys.exit(2)
    else:
        sys.stderr.write("[ERROR] Unsupported file type. Use .pdf, .docx, or .doc\n")
        sys.exit(2)


def _read_pdf_text(path: Path) -> str:
    """Extract text from a (non-scanned) PDF using pdfplumber, else pypdf."""
    # Try pdfplumber
    try:
        import pdfplumber  # type: ignore
        text_parts: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception:
        pass

    # Try pypdf (PyPDF2 or pypdf)
    try:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                text_parts.append("")
        return "\n".join(text_parts)
    except Exception:
        sys.stderr.write(
            "[WARN] Neither 'pdfplumber' nor 'pypdf' is available for PDF text extraction.\n"
            "Install one:\n  pip install pdfplumber\n  # or:\n  pip install pypdf\n"
        )
        return ""


def _read_pdf_via_ocr(path: Path) -> str:
    """OCR a scanned PDF via pdf2image + pytesseract. Requires Tesseract & Poppler installed."""
    try:
        import pytesseract  # type: ignore
    except Exception:
        sys.stderr.write(
            "[ERROR] Missing 'pytesseract' for OCR.\n"
            "Install it with:\n    pip install pytesseract\n"
            "Also ensure Tesseract is installed (macOS):\n    brew install tesseract\n"
        )
        sys.exit(2)
    try:
        from pdf2image import convert_from_path  # type: ignore
    except Exception:
        sys.stderr.write(
            "[ERROR] Missing 'pdf2image' for OCR.\n"
            "Install it with:\n    pip install pdf2image\n"
            "Also ensure Poppler is installed (macOS):\n    brew install poppler\n"
        )
        sys.exit(2)

    # Convert pages to images, then OCR each page
    try:
        images = convert_from_path(str(path))
    except Exception as e:
        sys.stderr.write(
            f"[ERROR] pdf2image failed to convert PDF (is Poppler installed?): {e}\n"
        )
        sys.exit(2)

    text_parts = []
    for img in images:
        try:
            text_parts.append(pytesseract.image_to_string(img))
        except Exception as e:
            sys.stderr.write(f"[WARN] OCR failed for a page: {e}\n")
            text_parts.append("")
    return "\n".join(text_parts)


# ----------------------------
# Redaction helpers
# ----------------------------

ORG_KEYWORDS = [
    "university", "college", "credit union", "bank", "trust", "llc", "inc",
    "corp", "company", "foundation", "department", "committee", "society"
]

# Built-in account name labels/brands/keywords (extensible via --account-term)
DEFAULT_ACCOUNT_TERMS = [
    "retirement account", "brokerage account", "checking account", "savings account",
    "money market", "credit card", "current account", "debit card", "investment account",
    "citibank", "fidelity", "tiaa-cref", "tiaa", "calpers", "santander", "hsbc",
    "bank of america", "wells fargo", "chase", "barclays", "capital one",
    "american express", "amex", "discover", "credit union"
]


def normalize_key(s: str) -> str:
    """ASCII lowercase, collapse whitespace and strip punctuation except hyphen/apostrophe for stable matching."""
    s_ascii = unidecode(s)
    s_ascii = s_ascii.replace("’", "'")
    s_ascii = re.sub(r"[^\w\s'\-\.]", " ", s_ascii)  # keep word chars, whitespace, apostrophe, hyphen, period
    s_ascii = re.sub(r"\s+", " ", s_ascii).strip()
    return s_ascii.lower()


def find_alias_links(text: str) -> Dict[str, str]:
    """
    Capture “Hereinafter ‘Alias’” style definitions and map alias -> full name (canonical string).
    Pattern rationale: simple, robust; avoids NLP; supports quotes “ ” ‘ ’ " ".
    """
    # Example: Dr. Jane A. Smith (hereinafter "Janie")
    title = r"(?:Dr|Mr|Mrs|Ms|Prof)\.\s+"
    token_word = r"(?:[A-Z][a-z]+(?:[-'][A-Za-z]+)*)"
    initial = r"(?:[A-Z]\.)"
    token = rf"(?:{token_word}|{initial})"
    full_name = rf"(?:{title})?{token}(?:\s+{token})+"  # ≥2 tokens
    # 'hereinafter' (optionally 'referred to as'), then quoted alias
    alias_pat = re.compile(
        rf"({full_name}).{{0,80}}?hereinafter(?:\s+referred\s+to\s+as)?\s*[“\"'‘]?([A-Za-z][\w\.\-']{{1,}})[”\"'’]?",
        re.IGNORECASE | re.DOTALL,
    )
    links: Dict[str, str] = {}
    for m in alias_pat.finditer(text):
        full = m.group(1)
        alias = m.group(2)
        # Strip leading titles from canonical full name in the mapping
        full_clean = re.sub(rf"^{title}", "", full).strip()
        links[normalize_key(alias)] = full_clean
    return links


def replace_person_names(
    text: str,
    include_allcaps: bool,
    mask_mode: bool,
    alias_links: Dict[str, str]
) -> Tuple[str, "OrderedDict[str, Dict[str, List[str]]]", int]:
    """
    Replace person names with deterministic placeholders or [REDACTED] and return:
      (new_text, placeholder_map, count_replacements)

    placeholder_map: OrderedDict mapping 'John Doe N' → {canonical: str, aliases: [..]}
    """
    # Build name pattern: optional title + ≥2 tokens; supports initials and hyphenated/apos names.
    title = r"(?:(?:Dr|Mr|Mrs|Ms|Prof)\.\s+)?"
    token_word = r"(?:[A-Z][a-z]+(?:[-'][A-Za-z]+)*)"
    initial = r"(?:[A-Z]\.)"
    token = rf"(?:{token_word}|{initial})"
    name_core = rf"{token}(?:\s+{token})+"  # at least two tokens
    # Optional possessive ’s or 's captured separately to preserve it.
    name_pat = re.compile(
        rf"(?<!\w){title}(?P<name>{name_core})(?P<pos>['’]s)?",
        re.MULTILINE
    )

    # State for deterministic numbering and mapping
    namekey_to_placeholder: "OrderedDict[str, str]" = OrderedDict()
    placeholder_map: "OrderedDict[str, Dict[str, List[str]]]" = OrderedDict()
    next_id = 1
    count = 0

    # convenience
    def assign_placeholder(canonical: str) -> str:
        nonlocal next_id
        k = normalize_key(canonical)
        if k not in namekey_to_placeholder:
            ph = f"John Doe {next_id}"
            namekey_to_placeholder[k] = ph
            placeholder_map[ph] = {"canonical": canonical, "aliases": []}
            next_id += 1
        return namekey_to_placeholder[k]

    # organization check
    def looks_like_org(s: str) -> bool:
        s_l = s.lower()
        return any(k in s_l for k in ORG_KEYWORDS)

    # ALL-CAPS detection (skip unless flag on)
    def is_all_caps_tokens(s: str) -> bool:
        words = re.findall(r"[A-Za-z][A-Za-z\-']*", s)
        words = [w for w in words if len(w) > 1]  # ignore 1-letter initials in caps check
        return bool(words) and all(w.isupper() for w in words)

    def repl(m: re.Match) -> str:
        nonlocal count
        raw = m.group(0)
        name_body = m.group("name")
        poss = m.group("pos") or ""

        # Idempotence: skip placeholders or [REDACTED]
        if re.match(r"^John Doe \d+\b", raw) or raw.strip().upper() == "[REDACTED]":
            return raw

        # Skip organizations
        if looks_like_org(raw):
            return raw

        # Skip ALL-CAPS unless requested (e.g., headings)
        if not include_allcaps and is_all_caps_tokens(name_body):
            return raw

        # Canonicalize: remove any leading title; preserve original spacing
        canonical = name_body.strip()

        # Alias linking: if matched name is an alias, map to its full canonical name
        alias_key = normalize_key(canonical)
        if alias_key in alias_links:
            canonical_full = alias_links[alias_key]
            placeholder = assign_placeholder(canonical_full)
            # track alias
            if normalize_key(canonical_full) != normalize_key(canonical):
                aliases = placeholder_map[placeholder]["aliases"]
                if canonical not in aliases:
                    aliases.append(canonical)
        else:
            placeholder = assign_placeholder(canonical)

        count += 1
        return ("[REDACTED]" if mask_mode else placeholder) + poss

    new_text = name_pat.sub(repl, text)
    return new_text, placeholder_map, count


def redact_addresses(text: str) -> Tuple[str, int]:
    """
    Redact address-like lines with [REDACTED ADDRESS].

    Heuristics:
      • Line matches: "<number> <street> <type>", e.g., "123 Main St", with common street types.
      • US "City, ST 12345" with optional ZIP+4.
      • UK postcodes (approximate but robust): "SW1A 1AA", "KY16 8QP" etc.
      • P.O. Box lines.
      • If a street line is redacted, redact the immediately following Apt/Apartment/Unit/Suite line too.
    """
    street_types = r"(?:St|Street|Rd|Road|Ave|Avenue|Blvd|Boulevard|Ln|Lane|Dr|Drive|Ct|Court|Broadway|Way|Ter|Terrace|Pl|Place|Pkwy|Parkway|Cir|Circle)\.?"
    street_line = re.compile(rf"^\s*\d{{1,6}}\s+[A-Za-z0-9\.\-\' ]+\s+{street_types}\s*$", re.IGNORECASE)
    po_box_line = re.compile(r"^\s*P\.?\s*O\.?\s*Box\s*\d+\s*$", re.IGNORECASE)
    city_state_zip = re.compile(r"^\s*[A-Za-z][A-Za-z\- ]+,\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?\s*$", re.IGNORECASE)
    uk_postcode_anywhere = re.compile(r"\b(?:[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2})\b", re.IGNORECASE)
    apt_next_line = re.compile(r"^\s*(?:Apt|Apartment|Unit|Suite|Ste)\s*[#\-]?\s*[A-Za-z0-9\-]+\s*$", re.IGNORECASE)

    lines = text.splitlines()
    redacted = 0
    i = 0
    out_lines: List[str] = []
    while i < len(lines):
        line = lines[i]
        is_street = bool(street_line.search(line)) or bool(po_box_line.search(line))
        is_city_zip = bool(city_state_zip.search(line))
        has_uk_postcode = bool(uk_postcode_anywhere.search(line))
        if is_street or is_city_zip or has_uk_postcode:
            out_lines.append("[REDACTED ADDRESS]")
            redacted += 1
            # If street, check next line for Apt/Apartment and redact it as well
            if is_street and (i + 1) < len(lines) and apt_next_line.search(lines[i + 1] or ""):
                out_lines.append("[REDACTED ADDRESS]")
                redacted += 1
                i += 2
                continue
            i += 1
            continue
        else:
            out_lines.append(line)
            i += 1
    return "\n".join(out_lines), redacted


def redact_account_names(text: str, custom_terms: List[str]) -> Tuple[str, int]:
    """
    Redact account labels/brands to [REDACTED ACCOUNT NAME].
    Terms are matched case-insensitively with word boundaries (hyphen allowed).
    """
    terms = [t.lower() for t in DEFAULT_ACCOUNT_TERMS] + [t.lower() for t in custom_terms]
    # Build a single regex with alternation. Allow spaces or hyphens between words inside the term.
    def term_to_pattern(t: str) -> str:
        parts = re.split(r"\s+", re.escape(t))
        return r"\b" + r"(?:\s+|-)?".join(parts) + r"\b"

    pattern = re.compile("|".join(term_to_pattern(t) for t in terms), re.IGNORECASE)
    count = 0

    def repl(m: re.Match) -> str:
        nonlocal count
        count += 1
        return "[REDACTED ACCOUNT NAME]"

    return pattern.sub(repl, text), count


def redact_account_numbers(text: str, strict_ids: bool) -> Tuple[str, int]:
    """
    Redact account numbers / IDs → [REDACTED ACCOUNT NUMBER].

    Patterns (applied in this order to avoid double-counting):
      1) account #<value>  (leave the 'account #' label)
      2) card-like #### #### #### #### (spaces/dashes optional)
      3) IBAN-like: 2 letters, 2 digits, 11–30 alnum
      4) 7+ digit runs
      5) VIN-like 11–17 uppercase A-HJ-NPR-Z0-9 (only with --strict-ids)
    """
    total = 0
    out = text

    # 1) account # XYZ
    acct_hash = re.compile(r"(?i)\b(account\s*#\s*)([A-Za-z0-9][A-Za-z0-9\-\s]{2,})\b")
    def repl_acct_hash(m: re.Match) -> str:
        nonlocal total
        total += 1
        return m.group(1) + "[REDACTED ACCOUNT NUMBER]"
    out = acct_hash.sub(repl_acct_hash, out)

    # 2) 16-digit card format (with optional spaces/dashes)
    card_like = re.compile(r"\b(?:\d{4}[ -]?){3}\d{4}\b")
    out, n = card_like.subn("[REDACTED ACCOUNT NUMBER]", out)
    total += n

    # 3) IBAN-like (very approximate; 15–34 total)
    iban_like = re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b")
    out, n = iban_like.subn("[REDACTED ACCOUNT NUMBER]", out)
    total += n

    # 4) 7+ contiguous digits (exclude those inside [REDACTED ...])
    long_digits = re.compile(r"(?<!\[REDACTED ACCOUNT NUMBER\])\b\d{7,}\b")
    out, n = long_digits.subn("[REDACTED ACCOUNT NUMBER]", out)
    total += n

    # 5) VIN-like (strict)
    if strict_ids:
        # 11–17 chars: exclude I,O,Q anywhere; uppercase letters or digits.
        vin_like = re.compile(r"\b[A-HJ-NPR-Z0-9]{11,17}\b")
        out, n = vin_like.subn("[REDACTED ACCOUNT NUMBER]", out)
        total += n

    return out, total


# ----------------------------
# Redaction pipeline
# ----------------------------

def redact_text_pipeline(
    text: str,
    include_allcaps: bool,
    mask_mode: bool,
    strict_ids: bool,
    extra_account_terms: List[str]
) -> RedactionResult:
    """Run the full redaction pipeline and return results with counts and mapping."""
    counts: Dict[str, int] = defaultdict(int)

    # 1) Link aliases first (before substitution)
    alias_links = find_alias_links(text)

    # 2) Person names
    text, placeholder_map, c = replace_person_names(
        text=text,
        include_allcaps=include_allcaps,
        mask_mode=mask_mode,
        alias_links=alias_links,
    )
    counts["persons"] += c

    # 3) Addresses (line-based)
    text, c = redact_addresses(text)
    counts["addresses"] += c

    # 4) Account names (labels/brands)
    text, c = redact_account_names(text, custom_terms=extra_account_terms)
    counts["account_names"] += c

    # 5) Account numbers / IDs
    text, c = redact_account_numbers(text, strict_ids=strict_ids)
    counts["account_numbers"] += c

    return RedactionResult(text=text, counts=dict(counts), placeholder_map=placeholder_map)


# ----------------------------
# Writers (DOCX & PDF)
# ----------------------------

def write_docx(path: Path, text: str, placeholder_map: "OrderedDict[str, Dict[str, List[str]]]", keep_key: bool) -> bool:
    """Write DOCX with paragraphs and optionally an anonymization key page."""
    try:
        import docx  # type: ignore
    except Exception:
        sys.stderr.write(
            "[WARN] Cannot write DOCX: 'python-docx' not installed.\n"
            "Install with:\n    pip install python-docx\n"
        )
        return False

    doc = docx.Document()
    # Main content
    for line in text.splitlines():
        doc.add_paragraph(line)

    # Key page
    if keep_key and placeholder_map:
        doc.add_page_break()
        doc.add_heading("Anonymization Key (Persons)", level=1)
        for ph, info in placeholder_map.items():
            aliases = info.get("aliases", [])
            aliastxt = ""
            if aliases:
                aliastxt = "  (aliases: " + ", ".join(sorted(set(aliases), key=lambda s: s.lower())) + ")"
            doc.add_paragraph(f"{ph} → {info.get('canonical', ph)}{aliastxt}")

    try:
        doc.save(str(path))
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to save DOCX: {e}\n")
        return False


def write_pdf(path: Path, text: str, placeholder_map: "OrderedDict[str, Dict[str, List[str]]]", keep_key: bool) -> bool:
    """Write simple PDF with Courier on A4 portrait and word-wrap; append key page."""
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        sys.stderr.write(
            "[WARN] Cannot write PDF: 'reportlab' not installed.\n"
            "Install with:\n    pip install reportlab\n"
        )
        return False

    W, H = A4
    margin = 50
    leading = 14
    font_name = "Courier"
    font_size = 11
    max_width = W - 2 * margin

    def draw_wrapped_lines(c: "canvas.Canvas", s: str) -> None:
        y = H - margin
        c.setFont(font_name, font_size)
        for raw_line in s.splitlines():
            # Simple word-wrap by measuring string width
            words = raw_line.split(" ")
            cur = ""
            for w in words:
                trial = (cur + " " + w).strip()
                if c.stringWidth(trial, font_name, font_size) <= max_width:
                    cur = trial
                else:
                    if y < margin + leading:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = H - margin
                    c.drawString(margin, y, cur)
                    y -= leading
                    cur = w
            if cur or raw_line == "":
                if y < margin + leading:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = H - margin
                c.drawString(margin, y, cur)
                y -= leading
        return

    try:
        c = canvas.Canvas(str(path), pagesize=A4)
        c.setAuthor("freedact.py")
        c.setTitle(path.name)
        draw_wrapped_lines(c, text)

        if keep_key and placeholder_map:
            c.showPage()
            c.setFont(font_name, font_size + 2)
            c.drawString(margin, H - margin, "Anonymization Key (Persons)")
            c.setFont(font_name, font_size)
            y = H - margin - 2 * leading
            for ph, info in placeholder_map.items():
                aliastxt = ""
                aliases = info.get("aliases", [])
                if aliases:
                    aliastxt = "  (aliases: " + ", ".join(sorted(set(aliases), key=lambda s: s.lower())) + ")"
                line = f"{ph} → {info.get('canonical', ph)}{aliastxt}"
                # Wrap key lines too
                for wrapped in textwrap.wrap(line, width=95):
                    if y < margin + leading:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = H - margin
                    c.drawString(margin, y, wrapped)
                    y -= leading

        c.save()
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to save PDF: {e}\n")
        return False


def write_json_key(path: Path, placeholder_map: "OrderedDict[str, Dict[str, List[str]]]") -> bool:
    """Write JSON mapping of placeholders to original names/aliases."""
    try:
        serializable = OrderedDict()
        for ph, info in placeholder_map.items():
            serializable[ph] = {
                "canonical": info.get("canonical", ""),
                "aliases": list(OrderedDict.fromkeys(info.get("aliases", [])))
            }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(serializable, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to write JSON key: {e}\n")
        return False


# ----------------------------
# Acceptance tests (--self-test)
# ----------------------------

def self_test() -> int:
    """
    Embedded acceptance tests (short strings). Asserts:
      1) Names & aliases (incl. possessive)
      2) Addresses
      3) Accounts (labels & numbers)
      4) Determinism (placeholder order)
      5) Mapping correctness (JSON key content)
    """
    # 1) Names & aliases with possessive
    sample = (
        'Dr. Jane A. Smith ... Hereinafter "Janie". Janie signed.\n'
        "Later, we saw Smith’s car parked outside.\n"
    )
    res = redact_text_pipeline(
        text=sample,
        include_allcaps=False,
        mask_mode=False,
        strict_ids=False,
        extra_account_terms=[],
    )
    assert "John Doe 1" in res.text, "Expected placeholder for person"
    assert "John Doe 1’s car" in res.text, "Possessive should be preserved"
    phs = list(res.placeholder_map.keys())
    assert phs and phs[0] == "John Doe 1", "First placeholder should be John Doe 1"
    info = res.placeholder_map["John Doe 1"]
    assert "Jane A. Smith" in info["canonical"], "Canonical full name recorded"
    assert any(a.lower() == "janie" for a in info["aliases"]), "Alias 'Janie' recorded"

    # 2) Addresses
    addr_sample = "123 Main St\nBoston, MA 02139\nPO Box 123\nSW1A 1AA\n"
    res2 = redact_text_pipeline(
        text=addr_sample,
        include_allcaps=False,
        mask_mode=False,
        strict_ids=False,
        extra_account_terms=[],
    )
    assert res2.text.strip().splitlines() == ["[REDACTED ADDRESS]"] * 4, "All address lines should be redacted"

    # 3) Accounts (labels & numbers)
    acct_sample = (
        "Checking Account 65355582\n"
        "account #ABC-123456\n"
        "GB29NWBK60161331926819\n"
        "4111 1111 1111 1111\n"
    )
    res3 = redact_text_pipeline(
        text=acct_sample,
        include_allcaps=False,
        mask_mode=False,
        strict_ids=False,
        extra_account_terms=[],
    )
    lines = res3.text.splitlines()
    assert "[REDACTED ACCOUNT NAME]" in lines[0], "Account label redacted"
    assert "[REDACTED ACCOUNT NUMBER]" in lines[0], "Account number redacted"
    assert "account #" in lines[1].lower() and "[REDACTED ACCOUNT NUMBER]" in lines[1], "Hash pattern redacted"
    assert "[REDACTED ACCOUNT NUMBER]" in lines[2], "IBAN redacted"
    assert "[REDACTED ACCOUNT NUMBER]" in lines[3], "Card number redacted"

    # 4) Determinism: order of placeholders by first appearance
    det_sample = "Dr. Ada Lovelace met Alan Turing. Then Ada spoke to Turing again."
    res4 = redact_text_pipeline(
        text=det_sample,
        include_allcaps=False,
        mask_mode=False,
        strict_ids=False,
        extra_account_terms=[],
    )
    # Ada first, then Alan
    ph_order = list(res4.placeholder_map.keys())
    assert ph_order == ["John Doe 1", "John Doe 2"], "Deterministic ordering of placeholders"

    # 5) Key page mapping exists and is correct
    assert res.placeholder_map["John Doe 1"]["canonical"].startswith("Jane"), "Key map canonical correct"

    print("All self-tests passed.")
    return 0


# ----------------------------
# Main program
# ----------------------------

def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()

    if args.self_test:
        code = self_test()
        sys.exit(code)

    if not args.input:
        parser.print_help(sys.stderr)
        sys.exit(2)

    in_path = Path(args.input)
    if not in_path.exists():
        sys.stderr.write(f"[ERROR] File not found: {in_path}\n")
        sys.exit(2)

    # Read text fully offline
    raw_text = read_input_text(in_path, use_ocr=args.ocr)

    # Normalize newlines
    raw_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")

    # Redact
    result = redact_text_pipeline(
        text=raw_text,
        include_allcaps=bool(args.include_allcaps),
        mask_mode=bool(args.mask_mode),
        strict_ids=bool(args.strict_ids),
        extra_account_terms=args.account_term or [],
    )

    # Summaries
    persons = result.counts.get("persons", 0)
    addresses = result.counts.get("addresses", 0)
    acct_names = result.counts.get("account_names", 0)
    acct_nums = result.counts.get("account_numbers", 0)
    total_placeholders = len(result.placeholder_map)

    # Report
    print("=== Redaction Summary ===")
    print(f" Persons replaced:        {persons}")
    print(f" Address lines redacted:  {addresses}")
    print(f" Account names redacted:  {acct_names}")
    print(f" Account numbers redacted:{acct_nums}")
    print(f" Unique persons (key):    {total_placeholders}")
    if args.dry_run:
        print("\n--dry-run: no files written.")
        return

    # Output paths
    base = in_path.with_suffix("")  # drop extension
    docx_out = Path(f"{base}_redacted.docx")
    pdf_out = Path(f"{base}_redacted.pdf")
    json_out = Path(f"{base}_redaction_key.json")

    # Write DOCX (always attempted)
    ok_docx = write_docx(docx_out, result.text, result.placeholder_map, keep_key=bool(args.keep_key))
    if ok_docx:
        print(f"Wrote DOCX: {docx_out}")
    else:
        print("Skipped DOCX (missing dependency).")

    # Write PDF if requested
    if args.pdf:
        ok_pdf = write_pdf(pdf_out, result.text, result.placeholder_map, keep_key=bool(args.keep_key))
        if ok_pdf:
            print(f"Wrote PDF:  {pdf_out}")
        else:
            print("Skipped PDF (missing dependency).")

    # Write JSON key (unless disabled)
    if args.keep_key:
        ok_json = write_json_key(json_out, result.placeholder_map)
        if ok_json:
            print(f"Wrote key:  {json_out}")
        else:
            print("Failed to write JSON key.")
    else:
        print("Key disabled (--no-keep-key): JSON/key page not written.")

    print("Done.")


if __name__ == "__main__":
    main()
