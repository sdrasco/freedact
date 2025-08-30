from __future__ import annotations

"""I/O helpers for Freedact.

This module groups functions that read text from supported document formats and
write redacted output files.  All operations here are offline and optional
dependencies are handled with helpful error messages.
"""

import io
import json
import re
import sys
import textwrap
from collections import OrderedDict
from pathlib import Path
from typing import Dict, List


# ----------------------------
# Reading inputs (offline)
# ----------------------------


def read_input_text(path: Path, use_ocr: bool) -> str:
    """
    Read the input file and return a single text string.
    Uses only offline libraries. Gracefully degrades with guidance if optional deps are missing.

    PDF: pdfplumber preferred, else pypdf. If both unavailable and --ocr is set, tries OCR via
         pdf2image + pytesseract (requires system installs).
    DOCX: python-docx
    DOC:  textract, else antiword (system), else instruct conversion.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf_text(path)
        if text and text.strip():
            return text
        if use_ocr:
            return _read_pdf_via_ocr(path)
        msg = (
            "[ERROR] No text extracted from PDF. Try installing 'pdfplumber' or 'pypdf'.\n"
            "       For scanned PDFs, re-run with --ocr and install:\n"
            "         pip install pytesseract pdf2image\n"
            "         brew install tesseract poppler\n"
        )
        sys.stderr.write(msg)
        sys.exit(2)

    elif suffix == ".docx":
        try:
            import docx  # type: ignore
        except Exception:
            sys.stderr.write(
                "[ERROR] Missing 'python-docx' for reading .docx.\n"
                "Install it with:\n    pip install python-docx\n"
            )
            sys.exit(2)
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    elif suffix == ".doc":
        # Try textract first
        try:
            import textract  # type: ignore
            return textract.process(str(path)).decode("utf-8", errors="ignore")
        except Exception:
            # Try antiword as a fallback
            try:
                from subprocess import check_output

                return check_output(["antiword", str(path)]).decode("utf-8", errors="ignore")
            except Exception:
                sys.stderr.write(
                    "[ERROR] Unable to read .doc file. Install 'textract' or 'antiword',"
                    " or convert the file to .docx/.pdf.\n"
                )
                sys.exit(2)

    else:
        sys.stderr.write(f"[ERROR] Unsupported file type: {suffix}\n")
        sys.exit(2)


def _read_pdf_text(path: Path) -> str:
    """Extract text from PDF using pdfplumber or pypdf."""
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages)
    except Exception:
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""


def _read_pdf_via_ocr(path: Path) -> str:
    """OCR PDFs when text extraction fails."""
    try:
        from pdf2image import convert_from_path  # type: ignore
    except Exception:
        sys.stderr.write(
            "[ERROR] Missing 'pdf2image' for OCR.\n"
            "Install it with:\n    pip install pdf2image\n"
            "Also ensure Poppler is installed (macOS):\n    brew install poppler\n"
        )
        sys.exit(2)

    try:
        import pytesseract  # type: ignore
    except Exception:
        sys.stderr.write(
            "[ERROR] Missing 'pytesseract' for OCR.\n"
            "Install it with:\n    pip install pytesseract\n"
            "Also ensure Tesseract OCR is installed.\n"
        )
        sys.exit(2)

    try:
        images = convert_from_path(str(path))
    except Exception as e:
        sys.stderr.write(
            f"[ERROR] pdf2image failed to convert PDF (is Poppler installed?): {e}\n"
        )
        sys.exit(2)

    text_parts = []
    for img in images:
        try:
            text_parts.append(pytesseract.image_to_string(img))
        except Exception as e:
            sys.stderr.write(f"[WARN] OCR failed for a page: {e}\n")
            text_parts.append("")
    return "\n".join(text_parts)


# ----------------------------
# Writers (DOCX & PDF)
# ----------------------------


def write_docx(
    path: Path,
    text: str,
    placeholder_map: "OrderedDict[str, Dict[str, List[str]]]",
    keep_key: bool,
) -> bool:
    """Write DOCX with paragraphs and optionally an anonymization key page."""
    try:
        import docx  # type: ignore
    except Exception:
        sys.stderr.write(
            "[WARN] Cannot write DOCX: 'python-docx' not installed.\n"
            "Install with:\n    pip install python-docx\n"
        )
        return False

    doc = docx.Document()
    for line in text.splitlines():
        doc.add_paragraph(line)

    if keep_key and placeholder_map:
        doc.add_page_break()
        doc.add_heading("Anonymization Key (Persons)", level=1)
        for ph, info in placeholder_map.items():
            aliases = info.get("aliases", [])
            aliastxt = ""
            if aliases:
                aliastxt = "  (aliases: " + ", ".join(sorted(set(aliases), key=lambda s: s.lower())) + ")"
            doc.add_paragraph(f"{ph} → {info.get('canonical', ph)}{aliastxt}")

    try:
        doc.save(str(path))
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to save DOCX: {e}\n")
        return False


def write_pdf(
    path: Path,
    text: str,
    placeholder_map: "OrderedDict[str, Dict[str, List[str]]]",
    keep_key: bool,
) -> bool:
    """Write simple PDF with Courier on A4 portrait and word-wrap; append key page."""
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        sys.stderr.write(
            "[WARN] Cannot write PDF: 'reportlab' not installed.\n"
            "Install with:\n    pip install reportlab\n"
        )
        return False

    W, H = A4
    margin = 50
    leading = 14
    font_name = "Courier"
    font_size = 11
    max_width = W - 2 * margin

    def draw_wrapped_lines(c: "canvas.Canvas", s: str) -> None:
        y = H - margin
        c.setFont(font_name, font_size)
        for raw_line in s.splitlines():
            words = raw_line.split(" ")
            cur = ""
            for w in words:
                trial = (cur + " " + w).strip()
                if c.stringWidth(trial, font_name, font_size) <= max_width:
                    cur = trial
                else:
                    if y < margin + leading:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = H - margin
                    c.drawString(margin, y, cur)
                    y -= leading
                    cur = w
            if cur or raw_line == "":
                if y < margin + leading:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = H - margin
                c.drawString(margin, y, cur)
                y -= leading

    try:
        c = canvas.Canvas(str(path), pagesize=A4)
        c.setAuthor("freedact.py")
        c.setTitle(path.name)
        draw_wrapped_lines(c, text)

        if keep_key and placeholder_map:
            c.showPage()
            c.setFont(font_name, font_size + 2)
            c.drawString(margin, H - margin, "Anonymization Key (Persons)")
            c.setFont(font_name, font_size)
            y = H - margin - 2 * leading
            for ph, info in placeholder_map.items():
                aliastxt = ""
                aliases = info.get("aliases", [])
                if aliases:
                    aliastxt = "  (aliases: " + ", ".join(sorted(set(aliases), key=lambda s: s.lower())) + ")"
                line = f"{ph} → {info.get('canonical', ph)}{aliastxt}"
                for wrapped in textwrap.wrap(line, width=95):
                    if y < margin + leading:
                        c.showPage()
                        c.setFont(font_name, font_size)
                        y = H - margin
                    c.drawString(margin, y, wrapped)
                    y -= leading

        c.save()
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to save PDF: {e}\n")
        return False


def write_json_key(path: Path, placeholder_map: "OrderedDict[str, Dict[str, List[str]]]") -> bool:
    """Write JSON mapping of placeholders to original names/aliases."""
    try:
        serializable = OrderedDict()
        for ph, info in placeholder_map.items():
            serializable[ph] = {
                "canonical": info.get("canonical", ""),
                "aliases": list(OrderedDict.fromkeys(info.get("aliases", [])))
            }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(serializable, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        sys.stderr.write(f"[ERROR] Failed to write JSON key: {e}\n")
        return False

